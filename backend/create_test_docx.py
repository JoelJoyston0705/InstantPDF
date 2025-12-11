from docx import Document

# Create a simple test DOCX file
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test paragraph to verify DOCX to PDF conversion.')
doc.add_paragraph('The conversion should work without Microsoft Word installed.')
doc.add_heading('Section 2', level=1)
doc.add_paragraph('Another paragraph with some content.')

doc.save('test_document.docx')
print("Test DOCX file created at test_document.docx")
